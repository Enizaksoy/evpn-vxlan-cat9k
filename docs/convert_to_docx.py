"""Convert the EVPN VXLAN HTML wiki document to a formatted Word (.docx) file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Define styles ──
style = doc.styles

# Normal
normal = style['Normal']
normal.font.name = 'Segoe UI'
normal.font.size = Pt(10)
normal.font.color.rgb = RGBColor(0x17, 0x2b, 0x4d)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

# Heading 1
h1 = style['Heading 1']
h1.font.name = 'Segoe UI'
h1.font.size = Pt(22)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x17, 0x2b, 0x4d)
h1.paragraph_format.space_before = Pt(0)
h1.paragraph_format.space_after = Pt(4)

# Heading 2
h2 = style['Heading 2']
h2.font.name = 'Segoe UI'
h2.font.size = Pt(16)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x17, 0x2b, 0x4d)
h2.paragraph_format.space_before = Pt(18)
h2.paragraph_format.space_after = Pt(8)

# Heading 3
h3 = style['Heading 3']
h3.font.name = 'Segoe UI'
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x17, 0x2b, 0x4d)
h3.paragraph_format.space_before = Pt(12)
h3.paragraph_format.space_after = Pt(4)

# Code style
code_style = style.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(8.5)
code_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.line_spacing = 1.1

# ── Helper functions ──

import re

# Syntax highlighting colors (terminal-style)
CLR_DEFAULT  = RGBColor(0x1E, 0x1E, 0x1E)  # near-black default
CLR_IP       = RGBColor(0x00, 0x6B, 0xB3)  # teal blue - IP addresses
CLR_MAC      = RGBColor(0xBF, 0x63, 0x00)  # amber - MAC addresses
CLR_GREEN    = RGBColor(0x00, 0x6B, 0x44)  # green - UP/FULL/Established
CLR_RED      = RGBColor(0xCC, 0x20, 0x00)  # red - DOWN/FAIL/DROPPED
CLR_PROMPT   = RGBColor(0x00, 0x80, 0x5A)  # green - prompts (sw1#, cumulus$)
CLR_IFACE    = RGBColor(0xAA, 0x55, 0x00)  # orange-brown - interface names
CLR_KEYWORD  = RGBColor(0x00, 0x52, 0xCC)  # blue - config keywords
CLR_VNI      = RGBColor(0x6C, 0x3F, 0xB5)  # purple - VNI/VLAN IDs
CLR_COMMENT  = RGBColor(0x6B, 0x77, 0x8C)  # gray - comments

# Compiled regex patterns for syntax highlighting
_SH_PATTERNS = [
    # Comments (# lines)
    (re.compile(r'^(\s*#.*)$'), CLR_COMMENT, False),
    # Prompts - sw1#, sw2#, cumulus@..., admin@...
    (re.compile(r'((?:sw\d+#|cumulus[@$][\w:~]*[$#>]?|admin@[\w\-]+[-cli>]*>?)\s?)'), CLR_PROMPT, True),
    # IP addresses with optional /mask
    (re.compile(r'\b(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}(?:/\d{1,2})?)\b'), CLR_IP, False),
    # MAC addresses colon format (aa:bb:cc:00:24:00)
    (re.compile(r'\b([0-9a-fA-F]{2}(?::[0-9a-fA-F]{2}){5})\b'), CLR_MAC, False),
    # MAC addresses dot format (aabb.cc00.2400)
    (re.compile(r'\b([0-9a-fA-F]{4}\.[0-9a-fA-F]{4}\.[0-9a-fA-F]{4})\b'), CLR_MAC, False),
    # Status positive
    (re.compile(r'\b(UP|FULL|Established|estab|FIXED|forwarding)\b', re.IGNORECASE), CLR_GREEN, True),
    # Status negative
    (re.compile(r'\b(DOWN|FAIL|ERROR|DROPPED|failed|blocked)\b', re.IGNORECASE), CLR_RED, True),
    # VNI numbers
    (re.compile(r'\b(401151|401152|401153|50101)\b'), CLR_VNI, False),
    # Interface names
    (re.compile(r'\b((?:Gi(?:gabitEthernet)?|GigabitEthernet|Te|nve|Loopback|Vlan|swp|vxlan|br_default|vlan)\d[\w/.\-]*)\b'), CLR_IFACE, False),
    # Versa interface names
    (re.compile(r'\b((?:tvi|vni|dtvi)-\d[\w/.\-]*)\b'), CLR_IFACE, False),
    # Config keywords (for config blocks)
    (re.compile(r'^\s*(router\s+(?:ospf|bgp)|interface|neighbor|address-family|l2vpn|vlan\s+configuration|member|nv\s+set)\b'), CLR_KEYWORD, True),
    # Versa config keywords
    (re.compile(r'^\s*(set\s+routing-instances|protocols|bridge-domains|bridge-options|policy-options)\b'), CLR_KEYWORD, True),
]


def colorize_line(line):
    """Split a line into (text, color, bold) segments for syntax highlighting."""
    # Check for full-line comment
    if line.strip().startswith('#') or line.strip().startswith('!'):
        return [(line, CLR_COMMENT, False)]

    # Find all matches with positions
    matches = []  # (start, end, color, bold)
    for pattern, color, bold in _SH_PATTERNS:
        for m in pattern.finditer(line):
            g = m.group(1) if m.lastindex else m.group(0)
            s = m.start(1) if m.lastindex else m.start(0)
            e = m.end(1) if m.lastindex else m.end(0)
            matches.append((s, e, color, bold))

    # Sort by position, remove overlaps (first match wins)
    matches.sort(key=lambda x: x[0])
    filtered = []
    last_end = 0
    for s, e, color, bold in matches:
        if s >= last_end:
            filtered.append((s, e, color, bold))
            last_end = e

    # Build segments
    segments = []
    pos = 0
    for s, e, color, bold in filtered:
        if s > pos:
            segments.append((line[pos:s], CLR_DEFAULT, False))
        segments.append((line[s:e], color, bold))
        pos = e
    if pos < len(line):
        segments.append((line[pos:], CLR_DEFAULT, False))

    if not segments:
        segments = [(line, CLR_DEFAULT, False)]

    return segments


def add_colored_runs(paragraph, line, font_name='Consolas', font_size=Pt(8.5)):
    """Add syntax-highlighted runs to a paragraph."""
    display_line = line if line.strip() else '\u00a0'
    segments = colorize_line(display_line)
    for text, color, bold in segments:
        run = paragraph.add_run(text)
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = color
        if bold:
            run.bold = True


def add_code_block(text, dark=False):
    """Add a code block with light bg + colored syntax, one paragraph per line for wiki compat."""
    # Strip HTML tags for clean text
    clean = re.sub(r'<[^>]+>', '', text)
    clean = clean.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
    clean = clean.replace('&harr;', '<->').replace('&mdash;', '—')
    clean = clean.replace('&rarr;', '->').replace('&larr;', '<-')
    # Escape square brackets to prevent Confluence from interpreting [ text ] as wiki macros
    clean = clean.replace('[', '[\u200B')
    clean = clean.strip()

    lines = clean.split('\n')
    first_p = None
    for i, line in enumerate(lines):
        p = doc.add_paragraph(style='CodeBlock')
        # Remove inter-line spacing so code block looks continuous
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if i == 0:
            p.paragraph_format.space_before = Pt(4)
            first_p = p
        if i == len(lines) - 1:
            p.paragraph_format.space_after = Pt(4)

        # Add light background shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F5F7" w:val="clear"/>')
        p.paragraph_format.element.get_or_add_pPr().append(shading)

        # Add syntax-highlighted runs
        add_colored_runs(p, line, font_size=Pt(8.5))
    return first_p


def add_panel(text, panel_type='info'):
    """Add a Confluence-style info/warning/error panel."""
    colors = {
        'info':    (RGBColor(0x00, 0x52, 0xcc), 'DEEBFF', 'Information'),
        'success': (RGBColor(0x00, 0x87, 0x5a), 'E3FCEF', 'Status'),
        'warning': (RGBColor(0xFF, 0x99, 0x1F), 'FFFAE6', 'Warning'),
        'error':   (RGBColor(0xDE, 0x35, 0x0B), 'FFEBE6', 'Alert'),
    }
    accent, bg, label = colors.get(panel_type, colors['info'])

    import re
    # Extract bold title if present
    bold_match = re.search(r'<strong>(.*?)</strong>', text)
    title = bold_match.group(1) if bold_match else label
    body = re.sub(r'<strong>.*?</strong>\s*', '', text)
    body = re.sub(r'<[^>]+>', '', body).replace('&lt;', '<').replace('&gt;', '>')
    body = body.replace('&mdash;', '—').replace('&harr;', '<->').strip()

    p = doc.add_paragraph()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    run_title = p.add_run(f'{title}: ')
    run_title.bold = True
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = accent

    run_body = p.add_run(body)
    run_body.font.name = 'Segoe UI'
    run_body.font.size = Pt(10)
    run_body.font.color.rgb = RGBColor(0x17, 0x2b, 0x4d)
    return p


def add_table(headers, rows, highlight_row=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Segoe UI'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x17, 0x2b, 0x4d)
        # Header shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F5F7" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Segoe UI'
            run.font.size = Pt(9)
            if highlight_row is not None and r_idx == highlight_row:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="DEEBFF" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                run.bold = True

    doc.add_paragraph()  # spacing
    return table


def add_code_to_cell(cell, text, font_size=Pt(7)):
    """Add syntax-highlighted code text into a table cell, one paragraph per line."""
    # Escape square brackets to prevent Confluence wiki macro interpretation
    text = text.replace('[', '[\u200B')

    # Clear existing paragraphs
    for p in cell.paragraphs:
        p.clear()

    lines = text.strip().split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        # Add syntax-highlighted colored runs
        add_colored_runs(p, line, font_size=font_size)


def add_comparison_table(title, cisco_text, versa_text, cumulus_text):
    """Add a 3-column comparison table with code in each cell."""
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Vendor headers with colors
    vendors = [
        ('Cisco Catalyst 9300 (sw2)', '0052CC'),
        ('Versa FlexVNF', 'FF6D00'),
        ('Cumulus Linux (NVIDIA)', '36B37E'),
    ]
    for i, (name, color) in enumerate(vendors):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.name = 'Segoe UI'
        run.font.size = Pt(9)
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F5F7" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Code cells
    texts = [cisco_text, versa_text, cumulus_text]
    for i, txt in enumerate(texts):
        add_code_to_cell(table.rows[1].cells[i], txt)

    doc.add_paragraph()  # spacing
    return table


def add_badge_text(paragraph, text, color_hex):
    """Add a badge-style inline text."""
    run = paragraph.add_run(f' [{text}] ')
    run.bold = True
    run.font.name = 'Segoe UI'
    run.font.size = Pt(8)
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)


# ══════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════

# ── Title Page ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
run = p.add_run('OnPrem Multi-Vendor\nEVPN VXLAN Integration')
run.font.name = 'Segoe UI'
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x17, 0x2b, 0x4d)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Cisco Catalyst 9300  •  Versa FlexVNF  •  Cumulus Linux (NVIDIA)')
run.font.name = 'Segoe UI'
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x6b, 0x77, 0x8c)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('Version 3.0  |  February 16, 2026  |  INTERNAL')
run.font.name = 'Segoe UI'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x6b, 0x77, 0x8c)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. Summary
# ══════════════════════════════════════════════════════════════════
doc.add_heading('1. Summary', level=2)

doc.add_paragraph(
    'This document covers the integration of a multi-vendor BGP EVPN VXLAN fabric for '
    'Layer 2 extension of VLAN 1151 (VNI 401151) across four leaf VTEPs. The fabric uses '
    'iBGP with a route reflector, OSPF underlay, and ingress replication for BUM traffic.'
)

# Status table
add_table(
    ['Component', 'Status'],
    [
        ['OSPF Adjacencies', 'FULL (all 4 peers)'],
        ['BGP EVPN Peers', 'Established (4/4)'],
        ['VXLAN Tunnels', '4 VTEPs UP'],
        ['Cross-vendor MAC Learning', 'Working'],
        ['VXLAN-GBP (Multi-Vendor)', 'FIXED'],
    ]
)

# ══════════════════════════════════════════════════════════════════
# 2. Topology
# ══════════════════════════════════════════════════════════════════
doc.add_heading('2. Topology', level=2)

topo_path = os.path.join(os.path.dirname(__file__), 'topology-for-docx.png')
if os.path.exists(topo_path):
    doc.add_picture(topo_path, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph('[Topology diagram - see diagrams/topology.svg]')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('All four leaf VTEPs connected to sw1 spine. VXLAN full-mesh tunnels overlay. Cisco (with SGT) and Versa send VXLAN-GBP headers.')
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x6b, 0x77, 0x8c)

# ══════════════════════════════════════════════════════════════════
# 3. Device Inventory
# ══════════════════════════════════════════════════════════════════
doc.add_heading('3. Device Inventory', level=2)

add_table(
    ['Device', 'Role', 'Platform', 'Loopback/VTEP', 'OSPF RID', 'BGP AS', 'Mgmt IP'],
    [
        ['[CISCO] sw1', 'Spine / BGP RR', 'Catalyst 9300', '1.1.1.1', '1.1.1.1', '6500', '192.168.20.77'],
        ['[CISCO] sw2', 'Leaf VTEP', 'Catalyst 9300', '1.1.1.2', '1.1.1.2', '6500', '192.168.20.76'],
        ['[CISCO] sw3', 'Leaf VTEP', 'Catalyst 9300', '1.1.1.3', '1.1.1.3', '6500', '192.168.20.81'],
        ['[VERSA] FlexVNF', 'SD-WAN VTEP', 'FlexVNF 22.1.4', '1.1.1.4', '12.12.12.10', '6500', '192.168.20.166'],
        ['[CUMULUS] Nvidia', 'Leaf VTEP', 'Cumulus Linux 5.x', '1.1.1.5', '1.1.1.5', '6500', '192.168.20.172'],
    ]
)

doc.add_heading('Underlay Point-to-Point Links', level=3)

add_table(
    ['Link', 'sw1 Intf', 'sw1 IP', 'Remote Intf', 'Remote IP', 'Subnet'],
    [
        ['sw1 <-> sw2', 'Gi1/0/1', '12.12.12.1', 'Gi1/0/1', '12.12.12.2', '12.12.12.0/30'],
        ['sw1 <-> sw3', 'Gi1/0/2', '12.12.12.5', 'Gi1/0/1', '12.12.12.6', '12.12.12.4/30'],
        ['sw1 <-> Versa', 'Gi1/0/3', '12.12.12.9', 'Vni-0/2', '12.12.12.10', '12.12.12.8/30'],
        ['sw1 <-> Cumulus', 'Gi1/0/4', '12.12.12.13', 'swp1', '12.12.12.14', '12.12.12.12/30'],
    ]
)

# ══════════════════════════════════════════════════════════════════
# 4. VLAN / VNI Mapping
# ══════════════════════════════════════════════════════════════════
doc.add_heading('4. VLAN / VNI Mapping', level=2)

add_table(
    ['VLAN', 'Name', 'VNI', 'Subnet', 'Purpose', 'Devices'],
    [
        ['1101', 'VRF1_CORE_VLAN', '50101', '—', 'L3VNI (Core Routing)', 'sw2, sw3'],
        ['1151', 'VRF1_ACCESS_VLAN', '401151', '192.168.1.0/25', 'L2 Extension (All VTEPs)', 'sw2, sw3, Versa, Cumulus'],
        ['1152', 'VRF1_ACCESS_VLAN_2', '401152', '192.168.2.0/25', 'Access VLAN', 'sw2, sw3'],
        ['1153', 'VRF1_ACCESS_VLAN_3', '401153', '192.168.3.0/25', 'Access VLAN', 'sw2, sw3'],
    ],
    highlight_row=1
)

# ══════════════════════════════════════════════════════════════════
# 5. Underlay Configuration (OSPF)
# ══════════════════════════════════════════════════════════════════
doc.add_heading('5. Underlay Configuration (OSPF)', level=2)

doc.add_heading('5.1 sw1 — Spine (Cisco IOS-XE)', level=3)

add_code_block("""! sw1 - OSPF underlay + interfaces
router ospf 100
 router-id 1.1.1.1
 auto-cost reference-bandwidth 100000

interface Loopback0
 ip address 1.1.1.1 255.255.255.255
 ip ospf 100 area 100

interface GigabitEthernet1/0/1
 description to sw2
 no switchport
 ip address 12.12.12.1 255.255.255.252
 ip ospf network point-to-point
 ip ospf 100 area 100

interface GigabitEthernet1/0/2
 description to sw3
 no switchport
 ip address 12.12.12.5 255.255.255.252
 ip ospf network point-to-point
 ip ospf 100 area 100

interface GigabitEthernet1/0/3
 description to Versa FlexVNF
 no switchport
 ip address 12.12.12.9 255.255.255.252
 ip mtu 1550
 ip ospf network point-to-point
 ip ospf mtu-ignore
 ip ospf 100 area 100

interface GigabitEthernet1/0/4
 description to Cumulus_Nvidia
 no switchport
 ip address 12.12.12.13 255.255.255.252
 ip mtu 1550
 ip ospf network point-to-point
 ip ospf mtu-ignore
 ip ospf 100 area 100""")

add_panel('MTU Mismatch: Cumulus defaults to MTU 9216. Cisco uses 1550. ip ospf mtu-ignore is configured on both ends to prevent OSPF ExStart/DBD failures.', 'warning')

doc.add_heading('5.2 Versa — Underlay OSPF', level=3)

add_code_block("""routing-instances routing-instance Underlay {
    instance-type virtual-router
    networks      [ OSPF_Underlay ]
    evpn-core
    interfaces    [ tvi-0/9005.0 ]

    policy-options {
        redistribution-policy Export-Vtep-To-OSPF {
            term VTEP_IP {
                match { address 1.1.1.4/32 }
                action { accept }
            }
        }
        redistribute-to-ospf Export-Vtep-To-OSPF
    }

    protocols {
        ospf 4023 {
            router-id     12.12.12.10
            area 100 {
                network-name OSPF_Underlay {
                    network-type    point-to-point
                    hello-interval  10
                    dead-interval   40
                }
            }
        }
    }
}""")

doc.add_heading('5.3 OSPF Verification (sw1)', level=3)

add_code_block("""sw1# show ip ospf neighbor

Neighbor ID     Pri   State           Dead Time   Address         Interface
12.12.12.10       0   FULL/  -        00:00:32    12.12.12.10     GigabitEthernet1/0/3
1.1.1.3           0   FULL/  -        00:00:39    12.12.12.6      GigabitEthernet1/0/2
1.1.1.2           0   FULL/  -        00:00:31    12.12.12.2      GigabitEthernet1/0/1
1.1.1.5           0   FULL/  -        00:00:37    12.12.12.14     GigabitEthernet1/0/4""", dark=False)

add_panel('OSPF Status: All four adjacencies FULL — sw2, sw3, Versa, Cumulus.', 'success')

# ══════════════════════════════════════════════════════════════════
# 6. BGP EVPN Overlay
# ══════════════════════════════════════════════════════════════════
doc.add_heading('6. BGP EVPN Overlay', level=2)

doc.add_heading('6.1 sw1 — Route Reflector', level=3)

add_code_block("""router bgp 6500
 bgp router-id interface Loopback0
 bgp log-neighbor-changes

 template peer-session spine-peer-session
  remote-as 6500
  update-source Loopback0
 exit-peer-session

 neighbor 1.1.1.2 inherit peer-session spine-peer-session
 neighbor 1.1.1.3 inherit peer-session spine-peer-session
 neighbor 1.1.1.4 inherit peer-session spine-peer-session
 neighbor 1.1.1.4 send-community both
 neighbor 1.1.1.5 inherit peer-session spine-peer-session
 neighbor 1.1.1.5 send-community both

 address-family l2vpn evpn
  neighbor 1.1.1.2 activate
  neighbor 1.1.1.2 send-community both
  neighbor 1.1.1.2 route-reflector-client
  neighbor 1.1.1.3 activate
  neighbor 1.1.1.3 send-community both
  neighbor 1.1.1.4 activate
  neighbor 1.1.1.4 send-community both
  neighbor 1.1.1.4 route-reflector-client
  neighbor 1.1.1.5 activate
  neighbor 1.1.1.5 send-community both
  neighbor 1.1.1.5 route-reflector-client
 exit-address-family""")

doc.add_heading('6.2 Versa — BGP EVPN', level=3)

add_code_block("""routing-instances routing-instance Underlay {
    protocols {
        bgp 3023 {
            router-id     1.1.1.4
            local-as { as-number 64515 }
            group vxlan_internal {
                type     internal
                family { l2vpn { evpn } }
                local-as 6500
                neighbor 1.1.1.1 {
                    local-address tvi-0/9005.0
                }
            }
        }
    }
}""")

add_panel('Note: Versa uses local-as 6500 override to match the fabric AS while keeping its internal AS 64515.', 'info')

doc.add_heading('6.3 BGP Verification (sw1)', level=3)

add_code_block("""sw1# show bgp l2vpn evpn summary
BGP router identifier 1.1.1.1, local AS number 6500

Neighbor        V    AS   MsgRcvd MsgSent  TblVer InQ OutQ Up/Down  State/PfxRcd
1.1.1.2         4  6500      224     290     106    0    0 03:03:30       18
1.1.1.3         4  6500      234     233     106    0    0 03:04:07       15
1.1.1.4         4  6500      429     454     106    0    0 03:03:42        1
1.1.1.5         4  6500     1148    1130       0    0    0 00:48:57        2""", dark=False)

add_panel('BGP Status: All four EVPN peers established. Cumulus: 2 routes (Type-2 + Type-3). Versa: 1 route (Type-3 IMET).', 'success')

# ══════════════════════════════════════════════════════════════════
# 7. L2VPN EVPN Instances
# ══════════════════════════════════════════════════════════════════
doc.add_heading('7. L2VPN EVPN Instances', level=2)

doc.add_heading('7.1 Route Target Mapping', level=3)

add_table(
    ['Device', 'RD', 'Export RT', 'Import RT'],
    [
        ['[CISCO] sw2', '1.1.1.2:1151 (auto)', '1151:1', '1151:1, 65000:1151, 2:2'],
        ['[CISCO] sw3', '1.1.1.3:1151 (auto)', '1151:1', '1151:1'],
        ['[VERSA]', '2L:122', '2:2, 1151:1', '2:2, 1151:1'],
        ['[CUMULUS]', '1.1.1.5:2 (auto)', '1151:1', '1151:1'],
    ]
)

add_panel('Route Target: Versa exports RT 2:2 in addition to 1151:1. All devices import 1151:1 so Versa routes are received.', 'warning')

doc.add_heading('7.2 sw2 L2VPN & NVE', level=3)

add_code_block("""l2vpn evpn instance 1151 vlan-based
 encapsulation vxlan
 route-target export 1151:1
 route-target import 1151:1
 route-target import 65000:1151
 route-target import 2:2
!
vlan configuration 1151
 member evpn-instance 1151 vni 401151
!
interface nve1
 no ip address
 source-interface Loopback0
 host-reachability protocol bgp
 member vni 50101 vrf VRF-1
 member vni 401151 ingress-replication
 member vni 401152 ingress-replication local-routing
 member vni 401153 ingress-replication local-routing""")

doc.add_heading('7.3 Versa Virtual-Switch', level=3)

add_code_block("""routing-instances routing-instance VERSA-default-switch {
    instance-type       virtual-switch
    bridge-options { l2vpn-service-type vlan }
    bridge-domains {
        VLAN-1151 {
            vlan-id 1151
            bridge-options { bridge-domain-arp-suppression enable }
            vxlan { vni 401151 }
        }
    }
    interfaces          [ vni-0/1.1 ]
    route-distinguisher 2L:122
    vrf-both-target     "target:2L:2 target:1151:1"
    protocols {
        evpn {
            vlan-id-list  [ 1151 ]
            encapsulation vxlan
            core-instance Underlay
        }
    }
}""")

# ══════════════════════════════════════════════════════════════════
# 8. Cumulus Linux (NVIDIA) — NVUE Configuration
# ══════════════════════════════════════════════════════════════════
doc.add_heading('8. Cumulus Linux (NVIDIA) — NVUE Configuration', level=2)

add_panel('Platform: Cumulus Linux 5.x uses NVUE (nv set) as the primary configuration method. Apply changes with nv config apply -y.', 'info')

doc.add_heading('8.1 Complete NVUE Commands', level=3)

add_code_block("""# -- Loopback (VTEP Source) --
nv set interface lo type loopback
nv set interface lo ip address 1.1.1.5/32

# -- Underlay Link to sw1 --
nv set interface swp1 type swp
nv set interface swp1 link state up
nv set interface swp1 ip address 12.12.12.14/30
nv set interface swp1 link mtu 9216

# -- OSPF Underlay --
nv set router ospf enable on
nv set router ospf router-id 1.1.1.5
nv set vrf default router ospf enable on
nv set vrf default router ospf router-id 1.1.1.5
nv set interface lo router ospf enable on
nv set interface lo router ospf area 100
nv set interface swp1 router ospf enable on
nv set interface swp1 router ospf area 100
nv set interface swp1 router ospf network-type point-to-point
nv set interface swp1 router ospf mtu-ignore on

# -- BGP EVPN Overlay --
nv set router bgp enable on
nv set router bgp autonomous-system 6500
nv set router bgp router-id 1.1.1.5
nv set vrf default router bgp enable on
nv set vrf default router bgp neighbor 1.1.1.1 remote-as internal
nv set vrf default router bgp neighbor 1.1.1.1 type numbered
nv set vrf default router bgp neighbor 1.1.1.1 update-source lo
nv set vrf default router bgp address-family l2vpn-evpn enable on
nv set vrf default router bgp neighbor 1.1.1.1 address-family l2vpn-evpn enable on

# -- VXLAN / NVE --
nv set nve vxlan enable on
nv set nve vxlan source address 1.1.1.5
nv set evpn enable on

# -- Bridge + VLAN 1151 + VNI --
nv set bridge domain br_default vlan 1151 vni 401151
nv set evpn vni 401151 route-target export 1151:1
nv set evpn vni 401151 route-target import 1151:1

# -- Trunk Port (swp2) - VLAN 1 native + VLAN 1151 tagged --
nv set interface swp2 type swp
nv set interface swp2 link state up
nv set interface swp2 bridge domain br_default
nv set interface swp2 bridge domain br_default stp admin-edge on
nv set interface swp2 bridge domain br_default stp bpdu-guard on
# NOTE: swp2 defaults to trunk mode (VLAN 1 PVID + all bridge VLANs tagged)
# Client uses tagged VLAN 1151 subinterface. Do NOT set access mode.

# -- SVI for VLAN 1151 --
nv set interface vlan1151 type svi
nv set interface vlan1151 vlan 1151
nv set interface vlan1151 ip address 192.168.1.15/25

# -- Apply --
nv config apply -y""")

doc.add_heading('8.2 Cumulus Verification', level=3)

add_code_block("""cumulus@cumulus:~$ sudo vtysh -c "show ip ospf neighbor"

Neighbor ID     Pri State           Dead Time Address         Interface
1.1.1.1           1 Full/DROther      39.892s 12.12.12.13     swp1:12.12.12.14""", dark=False)

add_code_block("""cumulus@cumulus:~$ sudo vtysh -c "show bgp l2vpn evpn summary"

BGP router identifier 1.1.1.5, local AS number 6500 vrf-id 0
Peers 1, using 23 KiB of memory

Neighbor        V   AS   MsgRcvd  MsgSent  Up/Down  State/PfxRcd  PfxSnt
1.1.1.1         4  6500    1148     1130   00:48:57           25       2""", dark=False)

add_code_block("""cumulus@cumulus:~$ sudo vtysh -c "show evpn vni 401151"

VNI: 401151
 Type: L2
 Vlan: 1151
 Bridge: br_default
 VxLAN interface: vxlan48
 Local VTEP IP: 1.1.1.5
 Remote VTEPs for this VNI:
  1.1.1.4 flood: HER
  1.1.1.3 flood: HER
  1.1.1.2 flood: HER
 Number of MACs (local and remote): 6
 Number of ARPs (local and remote): 6""", dark=False)

add_panel('Cumulus EVPN Status: VNI 401151 active. 3 remote VTEPs discovered. All remote MACs learned via EVPN Type-2 routes.', 'success')

# ══════════════════════════════════════════════════════════════════
# 9. NVE Peer Verification
# ══════════════════════════════════════════════════════════════════
doc.add_heading('9. NVE Peer Verification', level=2)

add_code_block("""sw2# show nve peers vni 401151

Interface  VNI      Type Peer-IP     RMAC/Num_RTs  eVNI     state  UP time
nve1       401151   L2CP 1.1.1.3     3             401151     UP   06:01:13
nve1       401151   L2CP 1.1.1.4     1             401151     UP   00:49:17
nve1       401151   L2CP 1.1.1.5     2             401151     UP   00:48:57""", dark=False)

add_code_block("""cumulus@cumulus:~$ ip -d link show vxlan48

14: vxlan48: <BROADCAST,MULTICAST,UP,LOWER_UP> mtu 9216
    master br_default state UNKNOWN
    vxlan id 401151 local 1.1.1.5 dstport 4789
    nolearning ttl 64 udpcsum gbp
    bridge_slave state forwarding learning off neigh_suppress on""", dark=False)

add_panel('Key Parameters: nolearning = EVPN controls FDB, neigh_suppress = ARP suppression, gbp = VXLAN-GBP enabled (required for interop with any GBP-sending VTEP: Cisco with SGT/GBP, Versa, etc.).', 'info')

# ══════════════════════════════════════════════════════════════════
# 10. MAC Address Learning
# ══════════════════════════════════════════════════════════════════
doc.add_heading('10. MAC Address Learning', level=2)

doc.add_heading('Versa Bridge Table', level=3)

add_code_block("""admin@Cisco-EVPN-Br1-cli> show bridge

ROUTING INSTANCE         BRIDGE     MAC-ADDRESS        TYPE  VLAN  TUNNEL EP
VERSA-default-switch     VLAN-1151  aa:bb:cc:00:27:00  CA    1151  1.1.1.3
VERSA-default-switch     VLAN-1151  aa:bb:cc:00:24:00  CA    1151  1.1.1.2""", dark=False)

doc.add_heading('sw2 EVPN MAC Table', level=3)

add_code_block("""sw2# show l2vpn evpn mac evi 1151

MAC Address    EVI   VLAN  Next Hop(s)
aabb.cc00.2400 1151  1151  Gi1/0/4:1151    (local)
aabb.cc00.2700 1151  1151  1.1.1.3         (remote - sw3)
aabb.cc00.2f10 1151  1151  1.1.1.4         (remote - Versa)""", dark=False)

add_panel('Cross-Vendor MAC Learning Confirmed: All VTEPs see each other\'s client MACs via EVPN Type-2 routes.', 'success')

# ══════════════════════════════════════════════════════════════════
# 11. VXLAN-GBP Interoperability Issue
# ══════════════════════════════════════════════════════════════════
doc.add_heading('11. VXLAN-GBP Interoperability Issue', level=2)

add_panel('<strong>Issue: NET-2026-001 — Severity: HIGH</strong> 100% data plane packet loss from any GBP-sending VTEP to Cumulus Linux. EVPN control plane is fully operational (BGP sessions up, routes exchanged, NVE peers visible). Root cause: Multiple vendors send VXLAN packets with GBP flags (0x88) — including Cisco Catalyst 9300 with group-based-policy + CTS/SGT enabled, and Versa FlexVNF (GBP always on by default). The Linux kernel VXLAN driver silently drops GBP-flagged packets on non-GBP VXLAN interfaces. This is not a vendor-specific issue — it is a Linux kernel behavior.', 'error')

doc.add_heading('11.1 What is VXLAN-GBP?', level=3)

doc.add_paragraph(
    'VXLAN-GBP (Group Based Policy) is an extension to the standard VXLAN header '
    '(draft-smith-vxlan-group-policy). It adds a Group Policy ID (SGT) to the VXLAN '
    'header for micro-segmentation.'
)

doc.add_paragraph('Standard VXLAN (RFC 7348): flags byte = 0x08 (I-bit only)')
doc.add_paragraph('VXLAN-GBP (Cisco SGT / Versa): flags byte = 0x88 (G-bit + I-bit)')

add_table(
    ['Bit', 'Name', 'Description'],
    [
        ['G (bit 0)', 'GBP Extension', 'When set, indicates GBP header present'],
        ['D (bit 3)', "Don't Learn", 'Tells remote VTEP not to learn source MAC'],
        ['I (bit 4)', 'VNI Valid', 'Same as standard VXLAN'],
        ['A (bit 7)', 'Applied', 'Policy already applied at source'],
        ['Group Policy ID', 'SGT', '16-bit Security Group Tag'],
    ]
)

doc.add_heading('11.2 Packet Capture Evidence', level=3)

doc.add_heading('Evidence A: Versa FlexVNF -> Cumulus (GBP always on)', level=4)
doc.add_paragraph('Captured on sw1 spine between Cumulus (1.1.1.5) and Versa (1.1.1.4):')

add_code_block("""Frame  Source (outer)           Destination (outer)      Flags   VNI
-----  ----------------------  -----------------------  ------  ------
11     1.1.1.5 Cumulus      -> 1.1.1.4 Versa            0x0800  401151
12     1.1.1.4 Versa        -> 1.1.1.5 Cumulus           0x8808  401151
13     1.1.1.5 Cumulus      -> 1.1.1.4 Versa            0x0800  401151
14     1.1.1.4 Versa        -> 1.1.1.5 Cumulus           0x8808  401151
...    (pattern repeats for ALL packets)""")

doc.add_heading('Evidence B: Cisco sw2 with SGT -> Cumulus (GBP + SGT 100)', level=4)
doc.add_paragraph('Captured on Cumulus swp1 with tcpdump -i swp1 -n \'udp port 4789\' -XX -vvv:')

add_code_block("""# Cisco sw2 (1.1.1.2, group-based-policy ON) -> Cumulus (1.1.1.5)
# ICMP Echo Reply: 192.168.1.22 -> 192.168.1.11

VXLAN Header (hex): 88 00 00 64 06 1e ff 00
                    ^^       ^^^^^
                    |        Group Policy ID = 0x0064 = SGT 100
                    Flags = 0x88 (G-bit + I-bit = GBP ENABLED)

# Cisco sw3 (1.1.1.3, NO group-based-policy) -> Cumulus (1.1.1.5)
# ICMP Echo Reply: 192.168.1.23 -> 192.168.1.11

VXLAN Header (hex): 08 00 00 00 06 1e ff 00
                    ^^
                    Flags = 0x08 (I-bit only = STANDARD VXLAN)""")

add_panel('<strong>Key Finding</strong> Cisco sw2 with group-based-policy + CTS/SGT sends identical GBP flags (0x88) as Versa FlexVNF. The only difference is the Group Policy ID value (Cisco inserts the configured SGT tag, e.g., 100). Both cause the same packet drops on Cumulus Linux without the gbp flag.', 'warning')

doc.add_heading('11.3 Raw Hex Dump Comparison', level=3)

add_code_block("""                    Flags Rsvd  GPID    VNI      Rsvd
                    ----- ----  ------  -------  ----
Cisco sw2 (GBP):   88    00    00 64   06 1e ff 00    <- SGT 100 in GPID field
Cisco sw3 (std):   08    00    00 00   06 1e ff 00    <- No GBP, GPID=0
Cumulus (std):     08    00    00 00   06 1e ff 00    <- No GBP, GPID=0
Versa (GBP):      88    00    xx xx   06 1e ff 00    <- GBP enabled, GPID varies""")

doc.add_heading('11.4 Why This Causes Failure', level=3)

doc.add_paragraph(
    'The Linux kernel VXLAN driver validates header flags strictly. Without the gbp option, '
    'the kernel expects flags = 0x08 (I-bit only). Receiving 0x88 (G-bit set) causes the '
    'packet to be silently dropped. No error, no log, no counter increment, no ICMP unreachable. '
    'This affects all GBP-sending VTEPs equally: Cisco with SGT, Versa, or any other vendor using VXLAN-GBP.'
)

add_code_block("""                                             GBP-Sending VTEP       Remote Client
Cumulus Client   Cumulus VTEP               (Cisco SGT / Versa)    192.168.1.x
192.168.1.12     1.1.1.5                    1.1.1.2 or 1.1.1.4
     |               |                          |                      |
     |-- ICMP Req -> |                          |                      |
     |               |== VXLAN flags=0x08 =====>|                      |
     |               |     (standard)           |-- ICMP Req -------> |
     |               |                          |                      |
     |               |                          |<-- ICMP Reply ------ |
     |               |<= VXLAN flags=0x88 ======|                      |
     |               |     (GBP!)               |                      |
     |            [DROPPED]                     |                      |
     |         kernel rejects GBP               |                      |
     |         on non-GBP iface                 |                      |""", dark=False)

doc.add_heading('11.5 Affected Traffic Flows', level=3)

add_table(
    ['Sender', 'Receiver', 'VXLAN Flags', 'Result'],
    [
        ['[CISCO] sw2 (GBP+SGT ON)', '[CUMULUS]', '0x88', 'DROPPED (without gbp flag)'],
        ['[VERSA] FlexVNF', '[CUMULUS]', '0x88', 'DROPPED (without gbp flag)'],
        ['[CISCO] sw3 (GBP OFF)', '[CUMULUS]', '0x08', 'OK'],
        ['[CUMULUS]', 'Any VTEP', '0x08', 'OK'],
        ['[CISCO] sw2 (GBP+SGT ON)', '[CISCO] sw3', '0x88', 'OK (IOS-XE ignores GBP)'],
    ]
)

add_panel('Cisco IOS-XE Receivers: Not Affected. IOS-XE silently ignores GBP bits and processes packets normally. This issue only affects Linux kernel VXLAN (Cumulus, Ubuntu, etc.).', 'info')

add_panel('<strong>Cisco as a Sender: Same Impact as Versa</strong> When group-based-policy is enabled under interface nve1 on Cisco Catalyst 9300, it sends VXLAN-GBP headers (flags=0x88) with SGT tag. This causes the exact same packet drops on Cumulus as Versa does. The issue is symmetric — any GBP sender triggers the Linux kernel drop.', 'warning')

# ══════════════════════════════════════════════════════════════════
# 12. VXLAN-GBP Fix & Resolution
# ══════════════════════════════════════════════════════════════════
doc.add_heading('12. VXLAN-GBP Fix & Resolution', level=2)

add_table(
    ['Attribute', 'Detail'],
    [
        ['Affected Traffic', 'Any GBP-sending VTEP -> Cumulus Linux (data plane only)'],
        ['GBP Senders', 'Cisco Catalyst 9300 (with group-based-policy + CTS/SGT), Versa FlexVNF (GBP always on)'],
        ['Root Cause', 'Linux kernel VXLAN driver drops packets with G-bit set (flags=0x88) on non-GBP interfaces. Not vendor-specific.'],
        ['Control Plane', 'Not affected — BGP EVPN routes exchanged correctly'],
        ['Cisco as Receiver', 'Not affected — IOS-XE ignores GBP bits from other senders'],
        ['Cisco as Sender', 'Causes same drop on Cumulus when SGT/GBP enabled (flags=0x88, SGT in GPID field)'],
        ['Fix Required On', 'Cumulus Linux (any Linux-based VTEP without gbp flag)'],
    ]
)

doc.add_heading('12.1 Step 1 — Recreate vxlan48 with GBP', level=3)

add_code_block("""# Delete existing
sudo ip link del vxlan48

# Recreate with GBP flag
sudo ip link add vxlan48 type vxlan id 401151 local 1.1.1.5 dstport 4789 nolearning gbp
sudo ip link set vxlan48 up""")

doc.add_heading('12.2 Step 2 — Restore bridge and VLAN settings', level=3)

add_code_block("""# Add back to bridge
sudo ip link set vxlan48 master br_default

# CRITICAL: Fix VLAN mapping (default PVID=1, must be 1151)
sudo bridge vlan del vid 1 dev vxlan48
sudo bridge vlan add vid 1151 dev vxlan48 pvid untagged

# Restore bridge port settings
sudo bridge link set dev vxlan48 state 3          # forwarding
sudo ip link set vxlan48 mtu 9216                 # match fabric MTU
sudo bridge link set dev vxlan48 learning off      # EVPN controls FDB
sudo bridge link set dev vxlan48 neigh_suppress on # ARP suppression
sudo bridge link set dev vxlan48 vlan_tunnel on""")

add_panel('VLAN Mapping Pitfall: When vxlan48 is recreated, the bridge assigns VLAN 1 as default PVID. This maps VNI 401151 to VLAN 1 instead of 1151, breaking ALL VXLAN traffic. Must explicitly fix.', 'error')

add_panel('Persistence: The manual ip link recreation is not persistent across reboots. NVUE does not currently expose a gbp option. Configure a post-boot script or edit /etc/network/interfaces.', 'warning')

doc.add_heading('12.3 Post-Fix Verification', level=3)

add_code_block("""cumulus@cumulus:~$ ip -d link show vxlan48 | grep -i gbp
    vxlan id 401151 local 1.1.1.5 dstport 4789 nolearning ttl 64 udpcsum gbp

cumulus@cumulus:~$ bridge vlan show dev vxlan48
port     vlan ids
vxlan48  1151 PVID Egress Untagged

cumulus@cumulus:~$ sudo vtysh -c "show evpn vni 401151" | grep Vlan
 Vlan: 1151          <- must be 1151, NOT 1""", dark=False)

add_panel('Resolution Confirmed: After enabling GBP, Cumulus accepts VXLAN-GBP packets from all GBP-sending VTEPs (Cisco with SGT and Versa). Full connectivity across all VTEPs.', 'success')

# ══════════════════════════════════════════════════════════════════
# 13. Command Reference
# ══════════════════════════════════════════════════════════════════
doc.add_heading('13. Command Reference', level=2)

doc.add_heading('Cisco IOS-XE', level=3)
add_table(
    ['Command', 'Purpose'],
    [
        ['show bgp l2vpn evpn summary', 'BGP EVPN neighbor status'],
        ['show bgp l2vpn evpn route-type 2', 'MAC/IP routes (Type-2)'],
        ['show bgp l2vpn evpn route-type 3', 'IMET routes (Type-3)'],
        ['show nve peers', 'VXLAN tunnel peer status'],
        ['show nve vni', 'VNI status / mappings'],
        ['show l2vpn evpn mac evi <id>', 'EVPN learned MACs per EVI'],
        ['show l2vpn evpn evi detail', 'Detailed EVI information'],
        ['show ip ospf neighbor', 'OSPF adjacency status'],
    ]
)

doc.add_heading('Cumulus Linux (NVIDIA)', level=3)
add_table(
    ['Command', 'Purpose'],
    [
        ['sudo vtysh -c "show bgp l2vpn evpn summary"', 'BGP EVPN peers'],
        ['sudo vtysh -c "show evpn vni <id>"', 'VNI status (VLAN map, remote VTEPs)'],
        ['sudo vtysh -c "show evpn mac vni <id>"', 'EVPN MAC table per VNI'],
        ['bridge fdb show dev vxlan48', 'Bridge FDB (MAC table)'],
        ['bridge vlan show dev vxlan48', 'VLAN membership check'],
        ['ip -d link show vxlan48', 'VXLAN interface detail (check GBP)'],
        ['nv config show --output commands', 'Running config in set format'],
    ]
)

doc.add_heading('Versa FlexVNF', level=3)
add_table(
    ['Command', 'Purpose'],
    [
        ['show bgp summary', 'BGP neighbor status'],
        ['show bridge', 'MAC address table'],
        ['show bridge-evpn', 'EVPN instance status'],
        ['show ospf route routing-instance Underlay', 'OSPF routes'],
    ]
)

# ══════════════════════════════════════════════════════════════════
# 14. Multi-Vendor Configuration Comparison (LANDSCAPE)
# ══════════════════════════════════════════════════════════════════
# Switch to landscape for wider comparison tables
landscape_section = doc.add_section()
landscape_section.orientation = WD_ORIENT.LANDSCAPE
new_w, new_h = landscape_section.page_height, landscape_section.page_width
landscape_section.page_width = new_w
landscape_section.page_height = new_h
landscape_section.top_margin = Cm(2)
landscape_section.bottom_margin = Cm(2)
landscape_section.left_margin = Cm(2)
landscape_section.right_margin = Cm(2)

doc.add_heading('14. Multi-Vendor Configuration Comparison', level=2)

doc.add_paragraph(
    'This section provides side-by-side comparison of configurations and verification '
    'outputs across all three vendor platforms for the same EVPN VXLAN fabric functions.'
)

# --- 14.1 OSPF Underlay ---
add_comparison_table(
    '14.1 OSPF Underlay Configuration',
    # Cisco (sw2)
    """router ospf 100
 router-id 1.1.1.2
 auto-cost reference-bandwidth 100000

interface Loopback0
 ip address 1.1.1.2 255.255.255.255
 ip ospf 100 area 100

interface GigabitEthernet1/0/1
 description to sw1
 no switchport
 ip address 12.12.12.2 255.255.255.252
 ip ospf network point-to-point
 ip ospf 100 area 100""",
    # Versa
    """set routing-instances Underlay
  instance-type virtual-router
  networks [ OSPF_Underlay ]
  interfaces [ tvi-0/9005.0 ]
  evpn-core

set routing-instances Underlay
  protocols ospf 4023
    router-id 12.12.12.10
    area 100
      network-name OSPF_Underlay
        network-type point-to-point
        hello-interval 10
        dead-interval 40

set routing-instances Underlay
  policy-options
    redistribution-policy Export-Vtep-To-OSPF
      term VTEP_IP
        match address 1.1.1.4/32
        action accept
    redistribute-to-ospf Export-Vtep-To-OSPF""",
    # Cumulus
    """# NVUE Commands:
nv set router ospf enable on
nv set router ospf router-id 1.1.1.5
nv set vrf default router ospf enable on
nv set vrf default router ospf router-id 1.1.1.5
nv set interface lo router ospf enable on
nv set interface lo router ospf area 100
nv set interface swp1 router ospf enable on
nv set interface swp1 router ospf area 100
nv set interface swp1 router ospf network-type
  point-to-point
nv set interface swp1 router ospf mtu-ignore on

# FRR Result:
router ospf
 ospf router-id 1.1.1.5
interface swp1
 ip ospf area 100
 ip ospf mtu-ignore
 ip ospf network point-to-point"""
)

# --- 14.2 BGP EVPN Overlay ---
add_comparison_table(
    '14.2 BGP EVPN Overlay Configuration',
    # Cisco (sw2)
    """router bgp 6500
 bgp router-id interface Loopback0
 bgp log-neighbor-changes

 neighbor 1.1.1.1 remote-as 6500
 neighbor 1.1.1.1 update-source Loopback0

 address-family l2vpn evpn
  neighbor 1.1.1.1 activate
  neighbor 1.1.1.1 send-community both
 exit-address-family""",
    # Versa
    """set routing-instances Underlay
  protocols bgp 3023
    router-id 1.1.1.4
    local-as as-number 64515
    group vxlan_internal
      type internal
      family l2vpn evpn
      local-as 6500
      neighbor 1.1.1.1
        local-address tvi-0/9005.0""",
    # Cumulus
    """# NVUE Commands:
nv set router bgp enable on
nv set router bgp autonomous-system 6500
nv set router bgp router-id 1.1.1.5
nv set vrf default router bgp enable on
nv set vrf default router bgp neighbor 1.1.1.1
  remote-as internal
nv set vrf default router bgp neighbor 1.1.1.1
  update-source lo
nv set vrf default router bgp address-family
  l2vpn-evpn enable on
nv set vrf default router bgp neighbor 1.1.1.1
  address-family l2vpn-evpn enable on

# FRR Result:
router bgp 6500
 bgp router-id 1.1.1.5
 neighbor 1.1.1.1 remote-as internal
 neighbor 1.1.1.1 update-source lo
 address-family l2vpn evpn
  neighbor 1.1.1.1 activate
  advertise-all-vni"""
)

# --- 14.3 VXLAN / L2VPN Instance ---
add_comparison_table(
    '14.3 VXLAN / L2VPN Instance Configuration',
    # Cisco (sw2)
    """l2vpn evpn instance 1151 vlan-based
 encapsulation vxlan
 route-target export 1151:1
 route-target import 1151:1
 route-target import 2:2

vlan configuration 1151
 member evpn-instance 1151 vni 401151

interface nve1
 no ip address
 source-interface Loopback0
 host-reachability protocol bgp
 member vni 401151 ingress-replication""",
    # Versa
    """set routing-instances VERSA-default-switch
  instance-type virtual-switch
  bridge-options
    l2vpn-service-type vlan
  bridge-domains VLAN-1151
    vlan-id 1151
    bridge-options
      bridge-domain-arp-suppression enable
    vxlan vni 401151
  interfaces [ vni-0/1.1 ]
  route-distinguisher 2L:122
  vrf-both-target
    "target:2L:2 target:1151:1"
  protocols evpn
    vlan-id-list [ 1151 ]
    encapsulation vxlan
    core-instance Underlay""",
    # Cumulus
    """# NVUE Commands:
nv set nve vxlan enable on
nv set nve vxlan source address 1.1.1.5
nv set evpn enable on
nv set bridge domain br_default
  vlan 1151 vni 401151
nv set evpn vni 401151
  route-target export 1151:1
nv set evpn vni 401151
  route-target import 1151:1

# /etc/network/interfaces Result:
auto vxlan48
iface vxlan48
  bridge-vlan-vni-map 1151=401151
  bridge-learning off

auto br_default
iface br_default
  bridge-ports swp2 vxlan48
  bridge-vlan-aware yes
  bridge-vids 1151"""
)

# --- 14.4 OSPF Verification ---
add_comparison_table(
    '14.4 OSPF Neighbor Verification',
    # Cisco (sw2)
    """sw2# show ip ospf neighbor

Neighbor ID  Pri State      Dead Time
             Address        Interface
1.1.1.1        0 FULL/  -   00:00:35
             12.12.12.1     Gi1/0/1""",
    # Versa
    """admin@Cisco-EVPN-Br1> show ospf

Router ID     Version Admin Op status Instance
12.12.12.10   2       en    up        4023

Intf address  Interface   State  Neighbor ID  Pri Op
12.12.12.9    vni-0/2.0   full   1.1.1.1      1   up""",
    # Cumulus
    """cumulus$ sudo vtysh -c "show ip ospf neighbor"

Neighbor ID  Pri State          Dead Time
             Address        Interface
1.1.1.1        1 Full/DROther   35.993s
             12.12.12.13    swp1:12.12.12.14"""
)

# --- 14.5 BGP EVPN Verification ---
add_comparison_table(
    '14.5 BGP EVPN Peer Verification',
    # Cisco (sw2)
    """sw2# show bgp l2vpn evpn summary

Neighbor  V    AS   MsgRcvd MsgSent
          TblVer InQ OutQ Up/Down  State
1.1.1.1   4  6500     844     816
          225    0    0 11:59:50    14
1.1.1.3   4  6500     822     822
          225    0    0 11:59:47    17""",
    # Versa
    """admin@Cisco-EVPN-Br1> show bgp summary

routing-instance: Underlay
BGP instance 3023
router-id : 1.1.1.4

l2vpn evpn statistics:
  iBGP routes in    : 25
  Active routes     : 23
  Advertised routes : 1

Neighbor   State    AS    Up/Down
1.1.1.1    estab    6500  12:00:01""",
    # Cumulus
    """cumulus$ sudo vtysh -c
  "show bgp l2vpn evpn summary"

BGP router identifier 1.1.1.5
local AS number 6500 vrf-id 0
Peers 1, using 23 KiB of memory

Neighbor  V    AS   MsgRcvd MsgSent
          Up/Down  State/PfxRcd PfxSnt
1.1.1.1   4  6500   13501   13291
          10:01:22           25     2"""
)

# --- 14.6 VXLAN / NVE Peer Verification ---
add_comparison_table(
    '14.6 VXLAN Tunnel / NVE Peer Verification',
    # Cisco (sw2)
    """sw2# show nve peers

Interface VNI    Type Peer-IP   eVNI   state UP time
nve1      401151 L2CP 1.1.1.3   401151   UP  00:43:10
nve1      401151 L2CP 1.1.1.4   401151   UP  00:40:58
nve1      401151 L2CP 1.1.1.5   401151   UP  00:43:10""",
    # Versa
    """admin@Cisco-EVPN-Br1> show bridge
  ingress-table routing-instance
  VERSA-default-switch

BDOMAIN    INTERFACE ENCAP OR VNI  TUNNEL
VLAN-1151  vni-0/1.1 N/A   N/A    N/A
           dtvi-0/79 VXLAN 401151 1.1.1.2
           dtvi-0/80 VXLAN 401151 1.1.1.3
           dtvi-0/81 VXLAN 401151 1.1.1.5""",
    # Cumulus
    """cumulus$ sudo vtysh -c
  "show evpn vni 401151"

VNI: 401151
 Type: L2
 Vlan: 1151
 VxLAN interface: vxlan48
 Local VTEP IP: 1.1.1.5
 Remote VTEPs for this VNI:
  1.1.1.4 flood: HER
  1.1.1.3 flood: HER
  1.1.1.2 flood: HER"""
)

# --- 14.7 MAC Address Table Verification ---
add_comparison_table(
    '14.7 MAC Address Table Verification (VLAN 1151)',
    # Cisco (sw2)
    """sw2# show l2vpn evpn mac evi 1151

MAC Address     EVI  VLAN  Next Hop(s)
5254.0091.2c07  1151 1151  1.1.1.3
5254.00c4.b4a1  1151 1151  Gi1/0/3:1151
aabb.cc00.2400  1151 1151  Gi1/0/4:1151
aabb.cc00.2700  1151 1151  1.1.1.3
aabb.cc00.3010  1151 1151  1.1.1.5""",
    # Versa
    """admin@Cisco-EVPN-Br1>
  show bridge mac-table brief

ROUTING INSTANCE     BRIDGE   INTERFACE
  MAC-ADDRESS        TYPE VLAN TUNNEL EP
VERSA-default-switch VLAN-1151 dtvi-0/79
  aa:bb:cc:00:24:00  CA   1151 1.1.1.2
VERSA-default-switch VLAN-1151 dtvi-0/80
  aa:bb:cc:00:27:00  CA   1151 1.1.1.3
VERSA-default-switch VLAN-1151 dtvi-0/81
  aa:bb:cc:00:30:10  CA   1151 1.1.1.5
VERSA-default-switch VLAN-1151 vni-0/1.1
  aa:bb:cc:00:2f:10  DA   1151 N/A""",
    # Cumulus
    """cumulus$ sudo vtysh -c
  "show evpn mac vni 401151"

MAC               Type  Intf/Remote
aa:bb:cc:00:27:00 remote 1.1.1.3
aa:bb:cc:00:30:10 local  vlan1151
aa:bb:cc:00:24:00 remote 1.1.1.2
aa:bb:cc:00:2f:10 remote 1.1.1.4
52:54:00:c4:b4:a1 remote 1.1.1.2
52:54:00:91:2c:07 remote 1.1.1.3"""
)

# Switch back to portrait for remaining content
portrait_section = doc.add_section()
portrait_section.orientation = WD_ORIENT.PORTRAIT
new_w, new_h = portrait_section.page_height, portrait_section.page_width
portrait_section.page_width = new_w
portrait_section.page_height = new_h
portrait_section.top_margin = Cm(2)
portrait_section.bottom_margin = Cm(2)
portrait_section.left_margin = Cm(2.5)
portrait_section.right_margin = Cm(2.5)

# ══════════════════════════════════════════════════════════════════
# 15. Lessons Learned
# ══════════════════════════════════════════════════════════════════
doc.add_heading('15. Lessons Learned', level=2)

add_table(
    ['#', 'Issue', 'Impact', 'Resolution'],
    [
        ['1', 'VXLAN-GBP Flag Mismatch', '100% data plane failure from any GBP-sending VTEP (Cisco with SGT, Versa) to Cumulus. Linux kernel VXLAN driver behavior, not vendor-specific.', 'Enable gbp on Cumulus VXLAN interface at creation time'],
        ['2', 'VLAN-VNI Mapping after Recreation', 'VNI maps to VLAN 1 instead of 1151', 'bridge vlan del vid 1, then add vid 1151 pvid untagged'],
        ['3', 'OSPF MTU Mismatch', 'OSPF stuck in ExStart (Cisco 1550 vs Cumulus 9216)', 'ip ospf mtu-ignore on both ends'],
        ['4', 'Cumulus NVUE OSPF Context', 'OSPF daemon won\'t activate', 'Must use: nv set vrf default router ospf enable on'],
        ['5', 'Versa Route Target', 'Versa exports RT 2:2 + 1151:1', 'All VTEPs import 1151:1. Optionally import 2:2.'],
        ['6', 'swp2 Access vs Trunk Mode', 'Client uses tagged VLAN 1151. Access port sends replies untagged, client drops them.', 'Keep swp2 as trunk: VLAN 1 native PVID + VLAN 1151 tagged. Do NOT use access mode.'],
    ]
)

doc.add_heading('Recommendations', level=3)

doc.add_heading('For Linux-based VTEPs (Cumulus, Ubuntu, etc.)', level=4)
bullets_linux = [
    'Always create VXLAN interfaces with gbp flag when integrating with any GBP-capable VTEP (Cisco with SGT, Versa, or others).',
    'Without gbp, the Linux kernel silently drops GBP-flagged packets — no logs, no counters, extremely difficult to diagnose.',
    'GBP Persistence: Edit /etc/network/interfaces or create post-boot script. NVUE does not support gbp natively.',
    'Monitoring: Check VXLAN flags after every nv config apply to ensure GBP preserved.',
]
for b in bullets_linux:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('For Cisco Catalyst 9300 with SGT/GBP', level=4)
bullets_cisco = [
    'Enabling group-based-policy on NVE changes the VXLAN header format from standard (0x08) to GBP-extended (0x88).',
    'This will break connectivity to any Linux-based VTEP that lacks the gbp interface flag.',
    'If SGT microsegmentation is not required, leave group-based-policy disabled to maintain standard VXLAN interoperability.',
]
for b in bullets_cisco:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('For Multi-Vendor Fabrics', level=4)
bullets_fabric = [
    'Test data plane connectivity after any GBP/SGT configuration change — control plane (BGP EVPN) will appear healthy even when data plane is completely broken.',
    'Packet Captures: Use tcpdump -i <intf> -n \'udp port 4789\' -XX to verify VXLAN flags byte (0x08 = standard, 0x88 = GBP).',
    'The issue is symmetric — Cisco with SGT causes the same drop on Cumulus as Versa does.',
    'MTU: Standardize underlay MTU across all vendors.',
]
for b in bullets_fabric:
    doc.add_paragraph(b, style='List Bullet')

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('INTERNAL ENGINEERING DOCUMENT')
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x6b, 0x77, 0x8c)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Multi-Vendor EVPN VXLAN — v3.0 — February 2026\nGenerated with Claude Code')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x6b, 0x77, 0x8c)

# ══════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════
base_name = os.path.join(os.path.dirname(__file__), 'Versa-Cisco-EVPN-VXLAN-Integration')
for suffix in ['_v3', '_v3a', '_v3b']:
    output_path = f'{base_name}{suffix}.docx'
    try:
        doc.save(output_path)
        break
    except PermissionError:
        continue
print(f'Word document saved: {output_path}')
print(f'Size: {os.path.getsize(output_path):,} bytes')
